import re
from docx import Document
from typing import List, Dict
from .config import settings
import os
from io import BytesIO
from .sharepoint_client import SharePointClient

# Reusable document-processing functions
PLACEHOLDER_PATTERN = re.compile(r"<([^<>]+)>")

def extract_placeholders_from_doc(doc: Document) -> List[str]:
    placeholders = set()
    for p in doc.paragraphs:
        text = "".join(run.text for run in p.runs)
        placeholders.update([m.strip() for m in PLACEHOLDER_PATTERN.findall(text)])
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = "".join(run.text for run in p.runs)
                    placeholders.update([m.strip() for m in PLACEHOLDER_PATTERN.findall(text)])
    return list(placeholders)

def extract_table_fields(doc: Document) -> List[str]:
    table_fields = []
    for table in doc.tables:
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        # detect schedule-style tables (example uses German days)
        if "montag" in headers and "dienstag" in headers:
            for r in range(1, len(table.rows)):
                row_label = table.rows[r].cells[0].text.strip()
                for c in range(1, len(headers)):
                    day = table.rows[0].cells[c].text.strip()
                    field = f"{row_label}_{day}"
                    table_fields.append(field)
    return table_fields

def extract_checkboxes(doc: Document) -> List[str]:
    checkboxes = []
    patterns = [r"☐\s*(.+)", r"\[\s?\]\s*(.+)"]

    def scan(text):
        for p in patterns:
            m = re.search(p, text)
            if m:
                opt = m.group(1).strip()
                if opt not in checkboxes:
                    checkboxes.append(opt)

    for p in doc.paragraphs:
        scan("".join(run.text for run in p.runs))

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    scan("".join(run.text for run in p.runs))

    return checkboxes

def fill_doc_instance(doc: Document, answers: Dict[str, str], selected_checkboxes: List[str], schedule: Dict[str, str]) -> None:
    def replace(paragraph):
        full = "".join(run.text for run in paragraph.runs)
        for key, val in answers.items():
            full = full.replace(f"<{key}>", val)
        for field, val in schedule.items():
            full = full.replace(field, val)
        for opt in selected_checkboxes:
            full = full.replace(f"☐ {opt}", f"☑ {opt}").replace(f"[ ] {opt}", f"[X] {opt}")
        if paragraph.runs:
            paragraph.runs[0].text = full
            for run in paragraph.runs[1:]:
                run.text = ""

    for p in doc.paragraphs:
        replace(p)

    for table in doc.tables:
        headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
        if "montag" in headers:
            for r in range(1, len(table.rows)):
                row_label = table.rows[r].cells[0].text.strip()
                for c in range(1, len(headers)):
                    day = table.rows[0].cells[c].text.strip()
                    key = f"{row_label}_{day}"
                    if key in schedule:
                        table.rows[r].cells[c].text = schedule[key]
        else:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        replace(p)

def generate_filled_document(template_name: str, answers: Dict[str,str], selected_checkboxes: List[str], schedule_values: Dict[str,str], sp_client: SharePointClient) -> str:
    # download template
    stream = sp_client.download_file(template_name)
    doc = Document(BytesIO(stream.read()))

    # fill
    fill_doc_instance(doc, answers, selected_checkboxes, schedule_values)

    # save locally
    os.makedirs(settings.LOCAL_OUTPUT_FOLDER, exist_ok=True)
    out_path = os.path.join(settings.LOCAL_OUTPUT_FOLDER, template_name.replace(".docx", "_Filled.docx"))
    doc.save(out_path)
    return out_path
